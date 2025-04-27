from docx import Document
from openpyxl import load_workbook
from datetime import datetime, timedelta
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def fill_invoice():

    try:
        excel_file = load_workbook("Sales Information.xlsx")
        excel_sheet = excel_file["Sales Information"]
    except Exception as e:
        print(f'Error loading Excel file: {e}')
        return

    start_row = 2

    while start_row <= excel_sheet.max_row:

        try:
            invoice_template = Document("Aurora Innovations Inc Invoice Template.docx")
        except Exception as e:
            print(f'Error loading invoice template: {e}')
            return

        row_index_list = []
        invoice_data = {}
        products_data = []
        subtotal = 0
        customer_id = None

        try:
            for row_index, row in enumerate(excel_sheet.iter_rows(min_row=start_row, values_only=True), start=start_row):
                if row[-1] == "FALSE":
                    customer_id = row[-3]
                    row_index_list.append(row_index)
                    break

            if customer_id is None:
                print("No valid customer ID found. Stopping execution.")
                return

            for row_index, row in enumerate(excel_sheet.iter_rows(min_row=start_row+1,
                                                                  values_only=True),
                                            start=start_row+1):
                if row[-3] == customer_id:
                    row_index_list.append(row_index)

            invoice_data["[Invoice Number]"] = (f'{excel_sheet[f'L{row_index_list[0]}'].value.strftime("%d%m%y")}-'
                                                    f'{excel_sheet[f'M{row_index_list[0]}'].value}')
            invoice_data["[Customer Name]"] = excel_sheet[f'A{row_index_list[0]}'].value
            invoice_data["[Company Name]"] = excel_sheet[f'B{row_index_list[0]}'].value
            invoice_data["[Customer Email]"] = excel_sheet[f'C{row_index_list[0]}'].value
            invoice_data["[Street Address]"] = excel_sheet[f'D{row_index_list[0]}'].value
            invoice_data["[City]"] = excel_sheet[f'E{row_index_list[0]}'].value
            invoice_data["[Credit Term]"] = excel_sheet[f'N{row_index_list[0]}'].value
            invoice_data["[Due Date]"] = (excel_sheet[f'L{row_index_list[0]}'].value + timedelta(days = int(excel_sheet[f'N{row_index_list[0]}'].value))).strftime("%d%m%y")

            file_output_name = f'invoices/{excel_sheet[f'B{row_index_list[0]}'].value}'

            for index in range(len(row_index_list)):
                subtotal += round(excel_sheet[f'I{row_index_list[index]}'].value, 2)

                products_data.append({
                    "Quantity": excel_sheet[f'H{row_index_list[index]}'].value,
                    "Description": excel_sheet[f'F{row_index_list[index]}'].value,
                    "Unit Price": round(excel_sheet[f'G{row_index_list[index]}'].value, 2),
                    "Line Total": round(excel_sheet[f'I{row_index_list[index]}'].value, 2)
                })
        except KeyError as e:
            print(f'Missing expected column in Excel: {e}')
            return
        except Exception as e:
            print(f'Error processing invoice data: {e}')
            return

        try:
            first_table = invoice_template.tables[0]
            second_table = invoice_template.tables[1]
            third_table = invoice_template.tables[2]
            fourth_table = invoice_template.tables[4]

            for table in [first_table, second_table, third_table, fourth_table]:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in invoice_data.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, str(value))

            for index in range(len(products_data)):
                new_row = third_table.add_row()
                new_row.cells[0].text = str(products_data[index].get("Quantity"))
                new_row.cells[1].text = str(products_data[index].get("Description"))
                new_row.cells[2].text = "{:.2f}".format(products_data[index].get("Unit Price"))
                new_row.cells[3].text = "{:.2f}".format(products_data[index].get("Line Total"))

                for cell in new_row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            sales_tax = round(0.09 * subtotal, 2)
            grand_total = round(subtotal + sales_tax, 2)

            subtotal_row = third_table.add_row()
            subtotal_row.cells[2].text = "SUBTOTAL"
            subtotal_row.cells[3].text = "{:.2f}".format(subtotal)

            sales_tax_row = third_table.add_row()
            sales_tax_row.cells[2].text = "SALES TAX"
            sales_tax_row.cells[3].text = "{:.2f}".format(sales_tax)

            grand_total_row = third_table.add_row()
            grand_total_row.cells[2].text = "GRAND TOTAL"
            grand_total_row.cells[3].text = "{:.2f}".format(grand_total)

            for row in [subtotal_row, sales_tax_row, grand_total_row]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                third_cell = row.cells[2]
                shading_elm = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))  # Yellow background
                third_cell._element.get_or_add_tcPr().append(shading_elm)

                for paragraph in third_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

            for cell in subtotal_row.cells[:2]:
                borders = parse_xml(r"""
                    <w:tcBorders {}>
                        <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                        </w:tcBorders>
                    """.format(nsdecls('w')))
                cell._element.get_or_add_tcPr().append(borders)

            for row in [subtotal_row, sales_tax_row, grand_total_row]:
                for index in [2, 3]:
                    cell = row.cells[index]
                    borders = parse_xml(r"""
                        <w:tcBorders {}>
                            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                            </w:tcBorders>
                        """.format(nsdecls('w')))
                    cell._element.get_or_add_tcPr().append(borders)


        except Exception as e:
            print(f'Error filling invoice template: {e}')
            return

        try:
            invoice_template.save(f'{file_output_name}.docx')
        except Exception as e:
            print(f'Error saving invoice file {file_output_name}: {e}')
            return

        try:
            for index in range(len(row_index_list)):
                excel_sheet[f'O{row_index_list[index]}'].value = 'TRUE'
            excel_file.save("Sales Information.xlsx")
        except Exception as e:
            print(f'Error updating Excel file: {e}')
            return

        start_row = max(row_index_list) + 1

    try:
        excel_file.close()
    except Exception as e:
        print(f'Error closing Excel file: {e}')
        return

fill_invoice()